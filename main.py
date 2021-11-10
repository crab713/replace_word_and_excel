from docx import Document
from xlrd import open_workbook # xlrd用于读取xld
from openpyxl import load_workbook
from xlutils.filter import process, XLRDReader, XLWTWriter
import os
import shutil
from win32com import client as wc


def copy2(wb):
    w = XLWTWriter()
    process(XLRDReader(wb, 'unknown.xls'), w)
    return w.output[0][1], w.style_list

def process_docx(file_dir, save_path, dic):
    if file_dir.split('\\')[-1][0] == '~':
        return
    try:
        if file_dir.split('.')[-1] == 'doc':
            word = wc.DispatchEx("Word.Application")
            doc = word.Documents.Open(file_dir)

            doc.SaveAs(file_dir+'x', 12)
            os.remove(file_dir)
            file_dir = file_dir+'x'
        doc = Document(file_dir)
    except:
        print('在处理doc格式文件%s时出现异常，转存docx失败'%file_dir)
        return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline = cell.paragraphs[0].runs
                    for i in inline:
                        for old in dic:
                            i.text = i.text.replace(old, dic[old])

    for paragraph in doc.paragraphs:
        if type(paragraph.text) == str:
            inline = paragraph.runs
            for i in inline:
                for old in dic:
                    i.text = i.text.replace(old, dic[old])
    print(save_path, '\\',file_dir.split('\\')[-1])
    doc.save(os.path.join(save_path,file_dir.split('\\')[-1]))
    doc.Close()

    

def process_xls(file_dir, save_path, dic):
    if file_dir.split('\\')[-1][0] == '~':
        return
    workbook = open_workbook(file_dir, formatting_info=True)  # 打开xls文件
    sheets= workbook.sheets() 
    workbook,s = copy2(workbook)
    for i in range(len(sheets)):
        sheet = sheets[i]
        work_sheet = workbook.get_sheet(i)
        nrows = sheet.nrows
        ncols = sheet.ncols
        for i in range(nrows):
            for j in range(ncols):
                data = sheet.cell(i,j).value
                if type(data) == str:
                    for old in dic:
                        data = data.replace(old, dic[old])
                    styles = s[sheet.cell_xf_index(i, j)]
                    work_sheet.write(i,j,data,styles)
    print(save_path, '\\',file_dir.split('\\')[-1])
    workbook.save(os.path.join(save_path,file_dir.split('\\')[-1]))

def process_xlsx(file_dir, save_path, dic):
    if file_dir.split('\\')[-1][0] == '~':
        return
    workbook = load_workbook(file_dir)
    sheets = workbook.worksheets
    for sheet in sheets:
        nrows = sheet.max_row
        ncols = sheet.max_column
        for i in range(nrows):
            for j in range(ncols):
                data = sheet.cell(i+1,j+1).value
                if type(data) == str:
                    for old in dic:
                        data = data.replace(old, dic[old])
                sheet.cell(i+1,j+1,data)
    print(save_path, '\\',file_dir.split('\\')[-1])
    workbook.save(os.path.join(save_path,file_dir.split('\\')[-1]))

def run(file_path, save_path, dic):
    for root,dirs,files in os.walk(file_path):
        for dir in dirs:
            save_dir = os.path.join(save_path, os.path.join(root, dir).replace(file_path+'\\', ''))
            if not os.path.exists(save_dir):
                os.makedirs(save_dir)
        folder = root.replace(file_path, '')
        for file in files:
            try:
                if file.split('.')[-1] == 'doc' or file.split('.')[-1] == 'docx':
                    process_docx(os.path.join(root, file), save_path+'\\'+folder, dic)
                elif file.split('.')[-1] == 'xls':
                    process_xls(os.path.join(root, file), save_path+'\\'+folder, dic)
                elif file.split('.')[-1] == 'xlsx':
                    process_xlsx(os.path.join(root, file), save_path+'\\'+folder, dic)
                else:
                    shutil.copy(os.path.join(root, file), os.path.join(save_path+'\\'+folder, file))     
            except:
                print('文件%s处理时出现异常，请手动修改或联系作者'%file)


# load args
file_path = input('输入需要处理的文件目录：\n')
print('开始录入替换文本，输入start结束输入，开始替换作业')
dic = dict()
old = ''
while old != 'start':
    old = input('输入需要替换的原字符串：')
    if old == 'start':
        break
    new = input('输入替换的新字符串：')
    dic[old] = new

print('开始作业')
# init
back_dir = os.path.dirname(os.path.abspath(file_path))
save_dir = os.path.join(back_dir, 'save')
if not os.path.exists(save_dir):
    os.makedirs(save_dir)

# start
run(file_path, save_dir, dic)
print('完成')
os.system('pause')


def forward(self, x):
    logit = self.model(x)
    label = argmax(x, dim=1)
    return logit, label